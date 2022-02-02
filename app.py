
from docx import Document
from docx.shared import Inches

document = Document
# Profile picture

document.add_picture('code pic.png', width=Inches(2.0))

# Name, Phone number and Email details

name = input('What is your full name? ')
phone_number = ('your phone numer? ')
email = ('your email? ')

document.add_paragraph(
      name + ' \n ' + phone_number + ' \n ' + email)

# About me
document.add_heading('about_me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

# Academic qualification
document.add_heading('Academic qualification')
p = document.add_paragraph()

School = input('Enter your school? ')
from_date = input('From date')
to_date = input('To date? ')

p.add_run(School + ' ').bold = True
p.add_run(from_date + '_' + to_date + '\n').italic = True

expirience_details = ('Describe your expirience at ' + School)
p.add_run(expirience_details)

# More Expirience or skill
while True:
      has_Skill= input('Do you have more Skill? Yes or No ')
      if has_Skill.lower() == 'yes':
            p = document.add_paragraph(Skill)

            Skill = input('Enter your Skill? ')
            from_date = input('From date')
            to_date = input('To date')
            p.add_run(Skill + ' ').bold = True
            p.add_run(from_date + '_' + to_date + '\n').italic = True
            Skill_details = ('Describe your expirience at ' + Skill)
            p.add_run(Skill)

      else:
            break


document.save('Abass_cv.docx')
 